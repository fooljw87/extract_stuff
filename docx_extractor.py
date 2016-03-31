from docx import Document
from openpyxl import load_workbook
import numpy as np
from pprint import pprint
import os

file_list = []
for file in os.listdir("C:\Word_extractor\word_files"):
    if file.endswith(".docx"):
        list.append(file_list, file)

wb = load_workbook('expenses.xlsx')
ws = wb['Form1']
#document_name = ('Document2.docx')
#document = Document(document_name)
#tables = document.tables
# column = table.columns.cells
# column = table.columns[1].cells

cost_dict = {}

def hasNumbers(inputString):
    return any(char.isdigit() for char in inputString)

def extract_document(document_name):
    document = Document("C:\Word_extractor\word_files\%s" %document_name)
    tables = document.tables
    cost_dict['%s' %document_name] = {}
    for table in tables:
        for column in table.columns:
            cells = column.cells
            for i in range(0,len(column.cells)):
                the_cell = cells[i].text.replace(' ','')
                if the_cell== '품명':
                    for add in range(0, (len(column.cells) - i)):
                        if cells[i+add].text != cells[i].text:
                            cost_dict['%s' %document_name]['품명'] = cells[i+add].text
                            break
                if the_cell == '단가':
                    for add in range(0, (len(column.cells) - i)):
                        if hasNumbers(cells[i+add].text) == True:
                            cost_dict['%s' %document_name]['단가'] = cells[i+add].text
                            break
                if the_cell == '수량':
                    for add in range(0, (len(column.cells) - i)):
                        if hasNumbers(cells[i+add].text) == True:
                            cost_dict['%s' %document_name]['수량'] = cells[i+add].text
                            break
                if the_cell == '단위':
                    for add in range(0, (len(column.cells) - i)):
                        if cells[i+add].text != cells[i].text:
                            cost_dict['%s' %document_name]['단위'] = cells[i+add].text
                            break
                if the_cell == '금액':
                    for add in range(0, (len(column.cells) - i)):
                        if cells[i+add].text != cells[i].text:
                            if hasNumbers(cells[i+add].text) == True:
                                cost_dict['%s' %document_name]['금액'] = cells[i+add].text
                                break

        for row in table.rows:
            cells = row.cells
            for i in range(0,len(row.cells)):
                the_cell = cells[i].text.replace(' ','')
                if the_cell == '최종Nego가격(부가세포함)':
                    for add in range(0, (len(row.cells) - i)):
                        if cells[i+add].text != cells[i].text:
                            if hasNumbers(cells[i+add].text) == True:
                                cost_dict['%s' %document_name]['최종금액'] = cells[i+add].text
                                break
                if the_cell == '최종Nego가격(VAT 별도)':
                    for add in range(0, (len(row.cells) - i)):
                        if cells[i+add].text != cells[i].text:
                            if hasNumbers(cells[i+add].text) == True:
                                cost_dict['%s' %document_name]['최종금액(VAT 별도)'] = cells[i+add].text
                                break

#extract_document('Document2.docx')
print(file_list)

for file in file_list:
    extract_document(file)

pprint(cost_dict)

xl_label = {'A': '품명',
            'B': '단가',
            'C' : '수량',
            'D': '단위',
            'E':'금액',
            'F':'최종금액'}

def export_to_xl():
    count = 1
    for thing in cost_dict:
        count += 1
        for val in xl_label.values():
            column = list(xl_label.keys())[list(xl_label.values()).index(val)]
            ws.cell('%s1' %column).value = xl_label[column]
            ws.cell('%s%s' %(column, count)).value = cost_dict[thing][xl_label[column]]

    wb.save("expenses.xlsx")

export_to_xl()
