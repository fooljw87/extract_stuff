from docx import Document
from openpyxl import load_workbook
import numpy as np
from pprint import pprint

wb = load_workbook('workbook.xlsx')
#document_name = ('Document2.docx')
#document = Document(document_name)
#tables = document.tables
# column = table.columns.cells
# column = table.columns[1].cells

cost_dict = {}

def extract_document(document_name):
    document = Document(document_name)
    tables = document.tables
    cost_dict['%s' %document_name] = {}
    for table in tables:
        target_row1 = None
        target_row2 = None
        target_row3 = None
        count = 0
        for row in table.rows:
            print('row')
#            print(row.cells[0].text)
            if count == 0:
                count += 1
                target_row1 = row
                cells1 = target_row1.cells
            elif count == 1:
                count += 1
                target_row2 = row
                cells2 = target_row2.cells
            elif count == 2:
                count -= 2
                target_row3 = row
                cells3 = target_row3.cells
            if target_row2 == None:
                continue
            if target_row3 == None:
                continue

            for i in range(0,len(target_row1.cells)):
                if cells1[i].text == '단가' or cells1[i].text =='단 가' or cells1[i].text =='단  가' or cells1[i].text =='단   가' or cells1[i].text =='단    가' or cells1[i].text =='단      가' or cells1[i].text =='단        가':
                    print(cells1[i].text)
                    print(cells2[i].text)
                    print(cells3[i].text)

#                    if cells2[i].text == int:
#                        cost_dict['%s' %document_name]['단가'] = cells2[i].text
#                    elif cells3[i].text == int:
#                        cost_dict['%s' %document_name]['단가'] = cells3[i].text


extract_document('Document2.docx')
print(cost_dict)
