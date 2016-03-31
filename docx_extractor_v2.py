from docx import Document
from openpyxl import load_workbook
import numpy as np
from pprint import pprint

wb = load_workbook('workbook.xlsx')
#document_name = ('Document2.docx')
#document = Document(document_name)
#tables = document.tables
# column = table.columns.cells
# column = table.columns[1].cells

cost_dict = {}

def hasNumbers(inputString):
    return any(char.isdigit() for char in inputString)

def extract_document(document_name):
    document = Document(document_name)
    tables = document.tables
    cost_dict['%s' %document_name] = {}
    for table in tables:
        for column in table.columns:
            len(column.cells)
            cells = column.cells
            for i in range(0,len(column.cells)):
                the_cell = cells[i].text.replace(' ','')
                if cells[i].text == '품명' or cells[i].text =='품 명' or cells[i].text =='품  명' or cells[i].text =='품   명' or cells[i].text =='품    명' or cells[i].text =='품      명' or cells[i].text =='품        명':
                    for add in range(0, (len(column.cells) - i)):
                        if cells[i+add].text != cells[i].text:
                            cost_dict['%s' %document_name]['품명'] = cells[i+add].text
                if cells[i].text == '단가' or cells[i].text =='단 가' or cells[i].text =='단  가' or cells[i].text =='단   가' or cells[i].text =='단    가' or cells[i].text =='단      가' or cells[i].text =='단        가':
                    for add in range(0, (len(column.cells) - i)):
                        if hasNumbers(cells[i+add].text) == True:
                            cost_dict['%s' %document_name]['단가'] = cells[i+add].text
                            break
                if cells[i].text == '수량' or cells[i].text =='수 량' or cells[i].text =='수  량' or cells[i].text =='수   량' or cells[i].text =='수    량' or cells[i].text =='수      량' or cells[i].text =='수        량':
                    for add in range(0, (len(column.cells) - i)):
                        if hasNumbers(cells[i+add].text) == True:
                            cost_dict['%s' %document_name]['수량'] = cells[i+add].text
                if cells[i].text == '단위' or cells[i].text =='단 위' or cells[i].text =='단  위' or cells[i].text =='단   위' or cells[i].text =='단    위' or cells[i].text =='단      위' or cells[i].text =='단        위':
                    for add in range(0, (len(column.cells) - i)):
                        if cells[i+add].text != cells[i].text:
                            cost_dict['%s' %document_name]['단위'] = cells[i+add].text
                if cells[i].text == '금액' or cells[i].text =='금 액' or cells[i].text =='금  액' or cells[i].text =='금   액' or cells[i].text =='금    액' or cells[i].text =='금      액' or cells[i].text =='금        액':
                    for add in range(0, (len(column.cells) - i)):
                        if hasNumbers(cells[i+add].text) == True:
                            cost_dict['%s' %document_name]['금액'] = cells[i+add].text
                            break
                    for add in range(0, (len(column.cells) - i)):
                        if hasNumbers(cells[i+add].text) == True:
                            cost_dict['%s' %document_name]['최종금액'] = cells[i+add].text



extract_document('Document2.docx')
print(cost_dict)
